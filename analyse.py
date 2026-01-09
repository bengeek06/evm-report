#!/usr/bin/env python3

import argparse
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor


def parser_arguments():
    """
    Parse les arguments de la ligne de commande
    """
    parser = argparse.ArgumentParser(
        description="Analyse EVM (Earned Value Management) - Analyse des dépenses et projections",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples d'utilisation:
  %(prog)s
  %(prog)s --sap EXPORT.xlsx --pv pv.xlsx --va va.xlsx --forecast forecast.xlsx
  %(prog)s --sap mes_depenses.xlsx --output graphique_evm.png
  %(prog)s --word rapport_evm.docx

Fichiers requis:
  - Fichier SAP : Export SAP des dépenses avec colonnes "Date de la pièce" et "Val./Devise objet"
  - Fichier PV  : Planned Value avec colonnes "Jalon", "Date", "Montant planifié"
  - Fichier VA  : Valeur Acquise avec colonnes "Jalon" et pourcentages mensuels
  - Fichier Forecast : Projections avec colonnes "Jalon", "Date projetée", "EAC (€)"

Rapport Word:
  L'option --word génère un rapport complet contenant définitions EVM, tableau, graphique
  et conclusion. Les fichiers intermédiaires (PNG, CSV, XLSX) sont automatiquement supprimés.
        """,
    )

    parser.add_argument(
        "--sap", default="EXPORT.XLSX", help="Fichier Excel d'export SAP des dépenses (défaut: EXPORT.XLSX)"
    )

    parser.add_argument("--pv", default="pv.xlsx", help="Fichier Excel de la Planned Value (défaut: pv.xlsx)")

    parser.add_argument("--va", default="va.xlsx", help="Fichier Excel de la Valeur Acquise (défaut: va.xlsx)")

    parser.add_argument(
        "--forecast", default="forecast.xlsx", help="Fichier Excel des projections (défaut: forecast.xlsx)"
    )

    parser.add_argument(
        "--output", default="analyse_evm.png", help="Fichier de sortie du graphique (défaut: analyse_evm.png)"
    )

    parser.add_argument(
        "--tableau", default="tableau_evm", help="Nom de base pour les fichiers tableau CSV/XLSX (défaut: tableau_evm)"
    )

    parser.add_argument("--word", help="Fichier Word de rapport complet à générer (ex: rapport_evm.docx)")

    parser.add_argument("--version", action="version", version="%(prog)s 1.0")

    return parser.parse_args()


def lire_export_sap(fichier="EXPORT.XLSX"):
    """
    Lit le fichier Excel d'export SAP des dépenses
    """
    try:
        df = pd.read_excel(fichier)
        print(f"Fichier chargé avec succès: {len(df)} lignes")
        print(f"Colonnes disponibles: {df.columns.tolist()}")
        return df
    except FileNotFoundError:
        print(f"Erreur: Le fichier {fichier} n'a pas été trouvé")
        return None
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier: {e}")
        return None


def lire_planned_value(fichier="pv.xlsx"):
    """
    Lit le fichier Excel contenant la Planned Value (PV)
    """
    try:
        df_pv = pd.read_excel(fichier)
        print(f"\nFichier PV chargé avec succès: {len(df_pv)} lignes")
        print(f"Colonnes PV: {df_pv.columns.tolist()}")
        return df_pv
    except FileNotFoundError:
        print(f"\n⚠️  Le fichier {fichier} n'a pas été trouvé")
        print("La Planned Value ne sera pas affichée.")
        return None
    except Exception as e:
        print(f"\n⚠️  Erreur lors de la lecture du fichier PV: {e}")
        return None


def lire_valeur_acquise(fichier="va.xlsx"):
    """
    Lit le fichier Excel contenant les pourcentages d'avancement (Valeur Acquise)
    """
    try:
        df_va = pd.read_excel(fichier)
        print(f"\nFichier VA chargé avec succès: {len(df_va)} lignes")
        print(f"Colonnes VA: {df_va.columns.tolist()}")
        return df_va
    except FileNotFoundError:
        print(f"\n⚠️  Le fichier {fichier} n'a pas été trouvé")
        print("La Valeur Acquise ne sera pas affichée.")
        return None
    except Exception as e:
        print(f"\n⚠️  Erreur lors de la lecture du fichier VA: {e}")
        return None


def lire_forecast(fichier="forecast.xlsx"):
    """
    Lit le fichier Excel contenant les projections (forecast)
    """
    try:
        df_forecast = pd.read_excel(fichier)
        print(f"\nFichier Forecast chargé avec succès: {len(df_forecast)} lignes")
        print(f"Colonnes Forecast: {df_forecast.columns.tolist()}")
        return df_forecast
    except FileNotFoundError:
        print(f"\n⚠️  Le fichier {fichier} n'a pas été trouvé")
        print("Les projections ne seront pas affichées.")
        return None
    except Exception as e:
        print(f"\n⚠️  Erreur lors de la lecture du fichier Forecast: {e}")
        return None


def calculer_depenses_cumulees(df, colonne_date="Date", colonne_montant="Montant"):
    """
    Calcule les dépenses cumulées mois par mois

    Args:
        df: DataFrame contenant les données
        colonne_date: nom de la colonne contenant les dates
        colonne_montant: nom de la colonne contenant les montants
    """
    # Copie du dataframe pour éviter les modifications
    df_work = df.copy()

    # Conversion de la colonne date en datetime si nécessaire
    if not pd.api.types.is_datetime64_any_dtype(df_work[colonne_date]):
        df_work[colonne_date] = pd.to_datetime(df_work[colonne_date])

    # Extraction du mois et année
    df_work["Mois"] = df_work[colonne_date].dt.to_period("M")

    # Agrégation par mois
    depenses_mensuelles = df_work.groupby("Mois")[colonne_montant].sum().sort_index()

    # Calcul des dépenses cumulées
    return depenses_mensuelles.cumsum()


def traiter_planned_value(df_pv):
    """
    Traite les données de Planned Value et calcule les montants cumulés
    """
    if df_pv is None:
        return None

    # Détection automatique des colonnes
    colonnes_date = ["Date", "date", "Date prévisionnelle", "Date prév.", "Mois"]
    colonnes_montant = [
        "Montant planifié",
        "Cumul planifié",
        "Montant",
        "montant",
        "Montant prévisionnel",
        "Montant prév.",
        "PV",
        "Planned Value",
    ]

    colonne_date = None
    colonne_montant = None

    for col in colonnes_date:
        if col in df_pv.columns:
            colonne_date = col
            break

    for col in colonnes_montant:
        if col in df_pv.columns:
            colonne_montant = col
            break

    if colonne_date is None or colonne_montant is None:
        print(f"\n⚠️  Colonnes PV non détectées. Colonnes disponibles: {df_pv.columns.tolist()}")
        return None

    print(f"✓ PV - Colonne date: {colonne_date}")
    print(f"✓ PV - Colonne montant: {colonne_montant}")

    df_work = df_pv.copy()

    # Conversion en datetime
    if not pd.api.types.is_datetime64_any_dtype(df_work[colonne_date]):
        df_work[colonne_date] = pd.to_datetime(df_work[colonne_date])

    # Extraction du mois
    df_work["Mois"] = df_work[colonne_date].dt.to_period("M")

    # Agrégation par mois
    pv_mensuelle = df_work.groupby("Mois")[colonne_montant].sum().sort_index()

    # Calcul cumulé
    pv_cumulee = pv_mensuelle.cumsum()

    # Interpolation linéaire pour remplir les mois manquants
    if len(pv_cumulee) > 1:
        # Créer une série avec tous les mois entre le premier et le dernier jalon
        premier_mois = pv_cumulee.index[0]
        dernier_mois = pv_cumulee.index[-1]

        # Générer tous les mois dans cet intervalle
        tous_les_mois = pd.period_range(start=premier_mois, end=dernier_mois, freq="M")

        # Réindexer avec tous les mois
        pv_cumulee_complete = pv_cumulee.reindex(tous_les_mois)

        # Interpolation linéaire pour les valeurs manquantes
        # Convertir en série temporelle pour l'interpolation
        pv_temp = pd.Series(pv_cumulee_complete.values, index=[p.to_timestamp() for p in pv_cumulee_complete.index])
        pv_temp = pv_temp.interpolate(method="time")

        # Reconvertir en Period index
        pv_cumulee = pd.Series(pv_temp.values, index=[pd.Period(d, freq="M") for d in pv_temp.index])

        print(f"✓ PV interpolée sur {len(pv_cumulee)} mois")

    # Récupération des jalons pour chaque mois
    jalons = {}
    if "Jalon" in df_pv.columns:
        for _, row in df_pv.iterrows():
            mois = pd.Period(row[colonne_date], freq="M")
            if mois not in jalons:
                jalons[mois] = []
            jalons[mois].append(row["Jalon"])

    return pv_cumulee, jalons


def calculer_earned_value(df_pv, df_va):
    """
    Calcule la Earned Value (EV) en combinant les montants planifiés et les pourcentages d'avancement
    """
    if df_pv is None or df_va is None:
        return None

    # Identification des colonnes de dates et montants dans PV
    colonnes_date = ["Date", "date", "Date prévisionnelle", "Date prév.", "Mois"]
    colonnes_montant = [
        "Montant planifié",
        "Cumul planifié",
        "Montant",
        "montant",
        "Montant prévisionnel",
        "Montant prév.",
        "PV",
        "Planned Value",
    ]

    colonne_date_pv = None
    colonne_montant_pv = None

    for col in colonnes_date:
        if col in df_pv.columns:
            colonne_date_pv = col
            break

    for col in colonnes_montant:
        if col in df_pv.columns:
            colonne_montant_pv = col
            break

    if colonne_date_pv is None or colonne_montant_pv is None:
        print("\n⚠️  Impossible de calculer l'EV: colonnes PV manquantes")
        return None

    # Identification de la colonne Jalon
    if "Jalon" not in df_pv.columns or "Jalon" not in df_va.columns:
        print("\n⚠️  Colonne 'Jalon' manquante dans PV ou VA")
        return None

    # Calcul de l'EV mois par mois
    # Les pourcentages dans VA sont cumulatifs par jalon
    # Pour chaque mois, on doit sommer l'EV de tous les jalons

    print("\n--- Debug calcul EV ---")

    # D'abord, collecter tous les mois présents
    tous_les_mois = set()
    for col in df_va.columns:
        try:
            date_col = pd.to_datetime(col)
            mois = date_col.to_period("M")
            tous_les_mois.add(mois)
        except Exception:
            continue

    tous_les_mois = sorted(tous_les_mois)

    # Pour chaque mois, calculer l'EV totale
    ev_par_mois = {}

    for mois_cible in tous_les_mois:
        ev_mois = 0

        # Parcourir chaque jalon
        for _idx, row_pv in df_pv.iterrows():
            jalon = row_pv["Jalon"]
            montant_jalon = row_pv[colonne_montant_pv]

            # Trouver le jalon correspondant dans VA
            row_va = df_va[df_va["Jalon"] == jalon]
            if row_va.empty:
                continue

            row_va = row_va.iloc[0]

            # Trouver le pourcentage d'avancement le plus récent pour ce mois
            pourcentage_actuel = 0
            for col in df_va.columns:
                try:
                    date_col = pd.to_datetime(col)
                    mois = date_col.to_period("M")

                    # Si ce mois est <= au mois cible, on prend le pourcentage
                    if mois <= mois_cible:
                        pourcentage = row_va[col]
                        if pd.notna(pourcentage):
                            pourcentage_actuel = max(pourcentage_actuel, pourcentage)
                except Exception:
                    continue

            # Calculer l'EV pour ce jalon (pourcentage est déjà en décimal: 1.0 = 100%)
            if pourcentage_actuel > 0:
                ev_jalon = montant_jalon * pourcentage_actuel
                ev_mois += ev_jalon

        if ev_mois > 0:
            ev_par_mois[mois_cible] = ev_mois

    # Afficher le détail pour debug
    for mois_cible in tous_les_mois:
        if mois_cible in ev_par_mois:
            print(f"\nMois {mois_cible}: EV totale = {ev_par_mois[mois_cible]:.2f}€")

            # Détail par jalon
            for _idx, row_pv in df_pv.iterrows():
                jalon = row_pv["Jalon"]
                montant_jalon = row_pv[colonne_montant_pv]

                row_va = df_va[df_va["Jalon"] == jalon]
                if row_va.empty:
                    continue

                row_va = row_va.iloc[0]

                pourcentage_actuel = 0
                for col in df_va.columns:
                    try:
                        date_col = pd.to_datetime(col)
                        mois = date_col.to_period("M")

                        if mois <= mois_cible:
                            pourcentage = row_va[col]
                            if pd.notna(pourcentage):
                                pourcentage_actuel = max(pourcentage_actuel, pourcentage)
                    except Exception:
                        continue

                if pourcentage_actuel > 0:
                    ev_jalon = montant_jalon * pourcentage_actuel
                    print(f"  {jalon}: {pourcentage_actuel*100:.1f}% -> {ev_jalon:.2f}€")

    if not ev_par_mois:
        print("\n⚠️  Aucune donnée EV calculée")
        return None

    # Convertir en Series et trier
    ev_series = pd.Series(ev_par_mois).sort_index()

    print(f"\n✓ Earned Value calculée pour {len(ev_series)} mois")

    return ev_series


def calculer_projections_automatiques(depenses_cumulees, ev_cumulee, pv_cumulee, _df_pv):
    """
    Calcule automatiquement les projections EAC selon différentes méthodes EVM
    """
    if ev_cumulee is None or len(ev_cumulee) == 0:
        return None

    print("\n--- Calcul des projections automatiques ---")

    # Valeurs actuelles
    ac_actuel = depenses_cumulees.iloc[-1]
    ev_actuel = ev_cumulee.iloc[-1]
    dernier_mois = depenses_cumulees.index[-1]

    # Budget à complétion (BAC)
    bac = pv_cumulee.iloc[-1] if pv_cumulee is not None else 0

    # Calcul des indices de performance
    cpi = ev_actuel / ac_actuel if ac_actuel > 0 else 1

    pv_actuel = pv_cumulee[dernier_mois] if pv_cumulee is not None and dernier_mois in pv_cumulee.index else 0
    spi = ev_actuel / pv_actuel if pv_actuel > 0 else 1

    # Calcul des EAC selon différentes méthodes
    projections = {}

    # Méthode 1 : Performance actuelle continue (CPI uniquement)
    eac_cpi = bac / cpi if cpi > 0 else bac
    projections["CPI"] = {
        "eac": eac_cpi,
        "formule": "BAC / CPI",
        "description": "Performance des coûts actuelle se poursuit",
        "reste": eac_cpi - ac_actuel,
    }

    # Méthode 2 : Performance coûts ET délais (CPI et SPI)
    eac_cpi_spi = ac_actuel + ((bac - ev_actuel) / (cpi * spi)) if (cpi * spi) > 0 else bac
    projections["CPI_SPI"] = {
        "eac": eac_cpi_spi,
        "formule": "AC + [(BAC-EV) / (CPI × SPI)]",
        "description": "Performance coûts ET délais se poursuit",
        "reste": eac_cpi_spi - ac_actuel,
    }

    # Méthode 3 : Reste comme prévu
    eac_reste_plan = ac_actuel + (bac - ev_actuel)
    projections["RESTE_PLAN"] = {
        "eac": eac_reste_plan,
        "formule": "AC + (BAC - EV)",
        "description": "Le reste se déroule comme prévu initialement",
        "reste": bac - ev_actuel,
    }

    # Estimation de la date de fin basée sur SPI
    if pv_cumulee is not None and spi > 0:
        dernier_mois_pv = pv_cumulee.index[-1]
        mois_restants_plan = (dernier_mois_pv.to_timestamp().year - dernier_mois.to_timestamp().year) * 12 + (
            dernier_mois_pv.to_timestamp().month - dernier_mois.to_timestamp().month
        )
        mois_restants_proj = int(mois_restants_plan / spi)
        date_fin_proj = dernier_mois + mois_restants_proj
    else:
        date_fin_proj = dernier_mois + 12  # Par défaut +12 mois

    # Générer les séries temporelles pour chaque projection
    mois_projection = pd.period_range(start=dernier_mois, end=date_fin_proj, freq="M")

    series_projections = {}
    for methode, data in projections.items():
        # Série avec AC historique + projection linéaire jusqu'à EAC
        serie = pd.Series(dtype=float)

        # Ajouter l'historique AC
        for mois in depenses_cumulees.index:
            serie[mois] = depenses_cumulees[mois]

        # Ajouter la projection linéaire
        if len(mois_projection) > 1:
            reste = data["reste"]
            nb_mois = len(mois_projection) - 1
            increment = reste / nb_mois if nb_mois > 0 else 0

            valeur_actuelle = ac_actuel
            for _i, mois in enumerate(mois_projection):
                if mois > dernier_mois:
                    valeur_actuelle += increment
                    serie[mois] = valeur_actuelle  # type: ignore[index]

        serie = serie.sort_index()
        series_projections[methode] = serie

        print(f"Méthode {methode}: EAC = {data['eac']:,.2f} € (Reste: {data['reste']:,.2f} €)")

    return projections, series_projections, date_fin_proj


def calculer_eac_projete(ev_cumulee, df_forecast, df_pv):
    """
    Calcule l'EAC projeté en combinant l'EV actuelle et les projections
    """
    if ev_cumulee is None or df_forecast is None or df_pv is None:
        return None

    print("\n--- Calcul EAC projeté ---")

    # Identifier les colonnes
    if "Date projetée" not in df_forecast.columns or "EAC (€)" not in df_forecast.columns:
        print("⚠️  Colonnes manquantes dans forecast.xlsx")
        return None

    # Créer un dictionnaire des montants EAC par jalon
    eac_par_jalon = {}
    for _, row in df_forecast.iterrows():
        jalon = row["Jalon"]
        date_projetee = row["Date projetée"]
        eac = row["EAC (€)"]

        # Convertir en period si nécessaire
        if isinstance(date_projetee, pd.Timestamp):
            mois_proj = date_projetee.to_period("M")
        else:
            mois_proj = pd.to_datetime(date_projetee).to_period("M")

        eac_par_jalon[jalon] = {"date": mois_proj, "eac": eac}

    # Récupérer tous les jalons avec leur montant planifié
    jalons_pv = {}
    for _, row in df_pv.iterrows():
        jalon = row["Jalon"]
        montant = row["Montant planifié"] if "Montant planifié" in df_pv.columns else 0
        jalons_pv[jalon] = montant

    # Créer une série temporelle pour l'EAC projeté
    # On commence avec l'EV actuelle
    dernier_mois_ev = ev_cumulee.index[-1]
    ev_finale = ev_cumulee.iloc[-1]

    # Collecter tous les mois de projection
    mois_projection = set()
    for info in eac_par_jalon.values():
        mois_projection.add(info["date"])

    mois_projection = sorted(mois_projection)

    # Calculer l'EAC cumulé mois par mois
    eac_par_mois = {}

    for mois in mois_projection:
        if mois <= dernier_mois_ev:
            # Pour les mois passés, on garde l'EV
            if mois in ev_cumulee.index:
                eac_par_mois[mois] = ev_cumulee[mois]
        else:
            # Pour les mois futurs, on accumule l'EAC
            eac_cumul = ev_finale
            for _jalon, info in eac_par_jalon.items():
                if info["date"] <= mois:
                    eac_cumul += info["eac"]
            eac_par_mois[mois] = eac_cumul
            print(f"Mois {mois}: EAC projeté = {eac_cumul:.2f}€")

    if not eac_par_mois:
        return None

    eac_series = pd.Series(eac_par_mois).sort_index()
    print(f"✓ EAC projeté calculé pour {len(eac_series)} mois")

    return eac_series, eac_par_jalon


def tracer_courbe(
    depenses_cumulees,
    pv_cumulee=None,
    jalons=None,
    ev_cumulee=None,
    eac_projete=None,
    jalons_forecast=None,
    fichier_sortie="analyse_evm.png",
):
    """
    Trace la courbe des dépenses cumulées, de la Planned Value et de l'Earned Value
    """
    plt.figure(figsize=(16, 9))

    # Conversion en k€ pour une meilleure lisibilité
    depenses_ke = depenses_cumulees / 1000

    # Conversion des périodes en dates pour l'affichage (dépenses réelles)
    dates_depenses = [periode.to_timestamp() for periode in depenses_cumulees.index]

    # Tracé des dépenses réelles (AC - Actual Cost)
    plt.plot(
        dates_depenses,
        depenses_ke.values,
        marker="o",
        linewidth=2.5,
        markersize=8,
        label="AC (Actual Cost - Dépenses réelles)",
        color="#e74c3c",
        zorder=3,
    )

    # Tracé de la Planned Value si disponible
    if pv_cumulee is not None:
        pv_ke = pv_cumulee / 1000
        dates_pv = [periode.to_timestamp() for periode in pv_cumulee.index]
        plt.plot(
            dates_pv,
            pv_ke.values,
            marker="s",
            linewidth=2.5,
            markersize=8,
            label="PV (Planned Value - Budget prévu)",
            color="#3498db",
            linestyle="--",
            zorder=2,
        )

        # Ajout des jalons sur la courbe PV
        if jalons:
            for periode, valeur in zip(pv_cumulee.index, pv_ke.values):
                if periode in jalons:
                    date = periode.to_timestamp()
                    jalons_str = "\n".join(jalons[periode])
                    plt.annotate(
                        jalons_str,
                        xy=(date, valeur),
                        xytext=(10, -15),
                        textcoords="offset points",
                        ha="left",
                        fontsize=8,
                        color="#3498db",
                        bbox={"boxstyle": "round,pad=0.3", "facecolor": "white", "edgecolor": "#3498db", "alpha": 0.8},
                    )

    # Tracé de l'Earned Value si disponible
    if ev_cumulee is not None:
        ev_ke = ev_cumulee / 1000
        dates_ev = [periode.to_timestamp() for periode in ev_cumulee.index]
        plt.plot(
            dates_ev,
            ev_ke.values,
            marker="^",
            linewidth=2.5,
            markersize=8,
            label="EV (Earned Value - Valeur acquise)",
            color="#2ecc71",
            linestyle="-.",
            zorder=2,
        )

    # Tracé de l'EAC projeté si disponible
    if eac_projete is not None:
        eac_ke = eac_projete / 1000
        dates_eac = [periode.to_timestamp() for periode in eac_projete.index]
        plt.plot(
            dates_eac,
            eac_ke.values,
            marker="*",
            linewidth=2.5,
            markersize=10,
            label="EAC (Estimate at Completion - Projection)",
            color="#f39c12",
            linestyle="--",
            zorder=2,
            alpha=0.8,
        )

        # Ajouter les jalons projetés sur la courbe EAC
        if jalons_forecast:
            for jalon, info in jalons_forecast.items():
                mois = info["date"]
                if mois in eac_projete.index:
                    date = mois.to_timestamp()
                    valeur = eac_projete[mois] / 1000
                    plt.annotate(
                        jalon,
                        xy=(date, valeur),
                        xytext=(10, 10),
                        textcoords="offset points",
                        ha="left",
                        fontsize=8,
                        color="#f39c12",
                        bbox={"boxstyle": "round,pad=0.3", "facecolor": "white", "edgecolor": "#f39c12", "alpha": 0.8},
                    )

    # Calcul et tracé de CV (Cost Variance) et SV (Schedule Variance)
    if ev_cumulee is not None and pv_cumulee is not None:
        # Trouver les mois communs pour CV
        mois_communs_cv = depenses_cumulees.index.intersection(ev_cumulee.index)
        if len(mois_communs_cv) > 0:
            cv_values = []
            cv_dates = []
            for mois in mois_communs_cv:
                cv = (ev_cumulee[mois] - depenses_cumulees[mois]) / 1000
                cv_values.append(cv)
                cv_dates.append(mois.to_timestamp())

            plt.plot(
                cv_dates,
                cv_values,
                marker="d",
                linewidth=2,
                markersize=6,
                label="CV (Cost Variance = EV - AC)",
                color="#e67e22",
                linestyle=":",
                zorder=1,
                alpha=0.7,
            )

            # Ajouter les valeurs de CV sur le dernier point
            if len(cv_dates) > 0:
                plt.annotate(
                    f"CV: {cv_values[-1]:.1f} k€",
                    xy=(cv_dates[-1], cv_values[-1]),
                    xytext=(10, 10),
                    textcoords="offset points",
                    ha="left",
                    fontsize=9,
                    color="#e67e22",
                    bbox={"boxstyle": "round,pad=0.4", "facecolor": "white", "edgecolor": "#e67e22", "alpha": 0.9},
                )

        # Trouver les mois communs pour SV
        mois_communs_sv = ev_cumulee.index.intersection(pv_cumulee.index)
        if len(mois_communs_sv) > 0:
            sv_values = []
            sv_dates = []
            for mois in mois_communs_sv:
                sv = (ev_cumulee[mois] - pv_cumulee[mois]) / 1000
                sv_values.append(sv)
                sv_dates.append(mois.to_timestamp())

            plt.plot(
                sv_dates,
                sv_values,
                marker="v",
                linewidth=2,
                markersize=6,
                label="SV (Schedule Variance = EV - PV)",
                color="#9b59b6",
                linestyle=":",
                zorder=1,
                alpha=0.7,
            )

            # Ajouter les valeurs de SV sur le dernier point
            if len(sv_dates) > 0:
                plt.annotate(
                    f"SV: {sv_values[-1]:.1f} k€",
                    xy=(sv_dates[-1], sv_values[-1]),
                    xytext=(10, -20),
                    textcoords="offset points",
                    ha="left",
                    fontsize=9,
                    color="#9b59b6",
                    bbox={"boxstyle": "round,pad=0.4", "facecolor": "white", "edgecolor": "#9b59b6", "alpha": 0.9},
                )

    plt.xlabel("Mois", fontsize=12)
    plt.ylabel("Montant cumulé (k€)", fontsize=12)
    plt.title("Analyse EVM - AC vs PV vs EV", fontsize=14, fontweight="bold")
    plt.grid(True, alpha=0.3)
    plt.legend(loc="upper left", fontsize=9)

    # Format des dates sur l'axe x
    plt.gcf().autofmt_xdate()

    # Ajout des valeurs sur les points (dépenses réelles) - uniquement quelques points clés
    step = max(1, len(dates_depenses) // 6)  # Afficher environ 6 labels
    for i, (date, valeur) in enumerate(zip(dates_depenses, depenses_ke.values)):
        if i % step == 0 or i == len(dates_depenses) - 1:
            plt.annotate(
                f"{valeur:.0f}",
                xy=(date, valeur),
                xytext=(0, 10),
                textcoords="offset points",
                ha="center",
                fontsize=7,
                color="#e74c3c",
            )

    plt.tight_layout()
    plt.savefig(fichier_sortie, dpi=300, bbox_inches="tight")
    print(f"Graphique sauvegardé: {fichier_sortie}")
    plt.show()


def tracer_courbe_realise(
    depenses_cumulees, pv_cumulee=None, jalons=None, ev_cumulee=None, fichier_sortie="analyse_evm_realise.png"
):
    """
    Trace le graphique du réalisé à date: AC, PV, EV avec CV et SV
    """
    plt.figure(figsize=(16, 9))

    # Conversion en k€ pour une meilleure lisibilité
    depenses_ke = depenses_cumulees / 1000

    # Conversion des périodes en dates pour l'affichage (dépenses réelles)
    dates_depenses = [periode.to_timestamp() for periode in depenses_cumulees.index]

    # Tracé des dépenses réelles (AC - Actual Cost)
    plt.plot(
        dates_depenses,
        depenses_ke.values,
        marker="o",
        linewidth=2.5,
        markersize=8,
        label="AC (Actual Cost - Dépenses réelles)",
        color="#e74c3c",
        zorder=3,
    )

    # Tracé de la Planned Value si disponible
    if pv_cumulee is not None:
        pv_ke = pv_cumulee / 1000
        dates_pv = [periode.to_timestamp() for periode in pv_cumulee.index]
        plt.plot(
            dates_pv,
            pv_ke.values,
            marker="s",
            linewidth=2.5,
            markersize=8,
            label="PV (Planned Value - Budget prévu)",
            color="#3498db",
            linestyle="--",
            zorder=2,
        )

        # Ajout des jalons sur la courbe PV
        if jalons:
            for periode, valeur in zip(pv_cumulee.index, pv_ke.values):
                if periode in jalons:
                    date = periode.to_timestamp()
                    jalons_str = "\n".join(jalons[periode])
                    plt.annotate(
                        jalons_str,
                        xy=(date, valeur),
                        xytext=(10, -15),
                        textcoords="offset points",
                        ha="left",
                        fontsize=8,
                        color="#3498db",
                        bbox={"boxstyle": "round,pad=0.3", "facecolor": "white", "edgecolor": "#3498db", "alpha": 0.8},
                    )

    # Tracé de l'Earned Value si disponible
    if ev_cumulee is not None:
        ev_ke = ev_cumulee / 1000
        dates_ev = [periode.to_timestamp() for periode in ev_cumulee.index]
        plt.plot(
            dates_ev,
            ev_ke.values,
            marker="^",
            linewidth=2.5,
            markersize=8,
            label="EV (Earned Value - Valeur acquise)",
            color="#2ecc71",
            linestyle="-.",
            zorder=2,
        )

    # Calcul et tracé de CV (Cost Variance) et SV (Schedule Variance)
    if ev_cumulee is not None and pv_cumulee is not None:
        # Trouver les mois communs pour CV
        mois_communs_cv = depenses_cumulees.index.intersection(ev_cumulee.index)
        if len(mois_communs_cv) > 0:
            cv_values = []
            cv_dates = []
            for mois in mois_communs_cv:
                cv = (ev_cumulee[mois] - depenses_cumulees[mois]) / 1000
                cv_values.append(cv)
                cv_dates.append(mois.to_timestamp())

            plt.plot(
                cv_dates,
                cv_values,
                marker="d",
                linewidth=2,
                markersize=6,
                label="CV (Cost Variance = EV - AC)",
                color="#e67e22",
                linestyle=":",
                zorder=1,
                alpha=0.7,
            )

            # Ajouter les valeurs de CV sur le dernier point
            if len(cv_dates) > 0:
                plt.annotate(
                    f"CV: {cv_values[-1]:.1f} k€",
                    xy=(cv_dates[-1], cv_values[-1]),
                    xytext=(10, 10),
                    textcoords="offset points",
                    ha="left",
                    fontsize=9,
                    color="#e67e22",
                    bbox={"boxstyle": "round,pad=0.4", "facecolor": "white", "edgecolor": "#e67e22", "alpha": 0.9},
                )

        # Trouver les mois communs pour SV
        mois_communs_sv = ev_cumulee.index.intersection(pv_cumulee.index)
        if len(mois_communs_sv) > 0:
            sv_values = []
            sv_dates = []
            for mois in mois_communs_sv:
                sv = (ev_cumulee[mois] - pv_cumulee[mois]) / 1000
                sv_values.append(sv)
                sv_dates.append(mois.to_timestamp())

            plt.plot(
                sv_dates,
                sv_values,
                marker="v",
                linewidth=2,
                markersize=6,
                label="SV (Schedule Variance = EV - PV)",
                color="#9b59b6",
                linestyle=":",
                zorder=1,
                alpha=0.7,
            )

            # Ajouter les valeurs de SV sur le dernier point
            if len(sv_dates) > 0:
                plt.annotate(
                    f"SV: {sv_values[-1]:.1f} k€",
                    xy=(sv_dates[-1], sv_values[-1]),
                    xytext=(10, -20),
                    textcoords="offset points",
                    ha="left",
                    fontsize=9,
                    color="#9b59b6",
                    bbox={"boxstyle": "round,pad=0.4", "facecolor": "white", "edgecolor": "#9b59b6", "alpha": 0.9},
                )

    plt.xlabel("Mois", fontsize=12)
    plt.ylabel("Montant cumulé (k€)", fontsize=12)
    plt.title("Analyse EVM - Réalisé à date (AC vs PV vs EV)", fontsize=14, fontweight="bold")
    plt.grid(True, alpha=0.3)
    plt.legend(loc="upper left", fontsize=9)

    # Format des dates sur l'axe x
    plt.gcf().autofmt_xdate()

    # Ajout des valeurs sur les points (dépenses réelles) - uniquement quelques points clés
    step = max(1, len(dates_depenses) // 6)  # Afficher environ 6 labels
    for i, (date, valeur) in enumerate(zip(dates_depenses, depenses_ke.values)):
        if i % step == 0 or i == len(dates_depenses) - 1:
            plt.annotate(
                f"{valeur:.0f}",
                xy=(date, valeur),
                xytext=(0, 10),
                textcoords="offset points",
                ha="center",
                fontsize=7,
                color="#e74c3c",
            )

    plt.tight_layout()
    plt.savefig(fichier_sortie, dpi=300, bbox_inches="tight")
    print(f"✓ Graphique du réalisé sauvegardé: {fichier_sortie}")


def tracer_courbe_projections(depenses_cumulees, ev_cumulee, projections, fichier_sortie="analyse_evm_projections.png"):
    """
    Trace le graphique des projections à terminaison: historique AC/EV + différents scénarios EAC

    Args:
        depenses_cumulees: Série temporelle des dépenses réelles (AC)
        ev_cumulee: Série temporelle de la valeur acquise (EV)
        projections: Dictionnaire contenant les projections calculées
                    {'cpi': {series, eac, date}, 'cpi_spi': {...}, 'reste_plan': {...}, 'forecast': {...}}
        fichier_sortie: Nom du fichier de sortie
    """
    plt.figure(figsize=(16, 9))

    # Conversion en k€ pour une meilleure lisibilité
    depenses_ke = depenses_cumulees / 1000
    ev_ke = ev_cumulee / 1000

    # Conversion des périodes en dates pour l'affichage
    dates_depenses = [periode.to_timestamp() for periode in depenses_cumulees.index]
    dates_ev = [periode.to_timestamp() for periode in ev_cumulee.index]

    # Tracé de l'historique AC et EV
    plt.plot(
        dates_depenses,
        depenses_ke.values,
        marker="o",
        linewidth=2.5,
        markersize=8,
        label="AC historique (Dépenses réelles)",
        color="#e74c3c",
        zorder=3,
    )

    plt.plot(
        dates_ev,
        ev_ke.values,
        marker="^",
        linewidth=2.5,
        markersize=8,
        label="EV historique (Valeur acquise)",
        color="#2ecc71",
        linestyle="-.",
        zorder=2,
    )

    # Couleurs et styles pour les projections
    couleurs = {"cpi": "#f39c12", "cpi_spi": "#9b59b6", "reste_plan": "#e67e22", "forecast": "#3498db"}

    labels_proj = {
        "reste_plan": "EAC reste à plan (optimiste)",
        "cpi": "EAC méthode CPI (réaliste)",
        "cpi_spi": "EAC méthode CPI×SPI (pessimiste)",
        "forecast": "EAC forecast manuel",
    }

    # Tracé de chaque projection
    for methode, couleur in couleurs.items():
        if methode in projections and projections[methode] is not None:
            proj_data = projections[methode]
            series = proj_data["series"] / 1000
            dates_proj = [periode.to_timestamp() for periode in series.index]
            eac_final = proj_data["eac"] / 1000
            date_fin = proj_data["date"]

            plt.plot(
                dates_proj,
                series.values,
                marker="*",
                linewidth=2.5,
                markersize=8,
                label=f"{labels_proj[methode]} ({eac_final:.0f} k€)",
                color=couleur,
                linestyle="--",
                zorder=2,
                alpha=0.8,
            )

            # Ajouter annotation sur le point final
            if len(dates_proj) > 0:
                plt.annotate(
                    f'{eac_final:.0f} k€\n{date_fin.strftime("%m/%Y")}',
                    xy=(dates_proj[-1], series.to_numpy()[-1]),
                    xytext=(10, 10),
                    textcoords="offset points",
                    ha="left",
                    fontsize=8,
                    color=couleur,
                    bbox={"boxstyle": "round,pad=0.3", "facecolor": "white", "edgecolor": couleur, "alpha": 0.8},
                )

    plt.xlabel("Mois", fontsize=12)
    plt.ylabel("Montant cumulé (k€)", fontsize=12)
    plt.title("Analyse EVM - Projections à terminaison (scénarios EAC)", fontsize=14, fontweight="bold")
    plt.grid(True, alpha=0.3)
    plt.legend(loc="upper left", fontsize=9)

    # Format des dates sur l'axe x
    plt.gcf().autofmt_xdate()

    plt.tight_layout()
    plt.savefig(fichier_sortie, dpi=300, bbox_inches="tight")
    print(f"✓ Graphique des projections sauvegardé: {fichier_sortie}")


def generer_tableau_comparatif(depenses_cumulees, pv_cumulee=None, ev_cumulee=None, nom_base="tableau_evm"):
    """
    Génère un tableau comparatif des valeurs PV, AC et EV
    """
    # Créer un DataFrame avec toutes les périodes
    toutes_periodes = set(depenses_cumulees.index)
    if pv_cumulee is not None:
        toutes_periodes.update(pv_cumulee.index)
    if ev_cumulee is not None:
        toutes_periodes.update(ev_cumulee.index)

    toutes_periodes = sorted(toutes_periodes)

    # Créer le tableau
    data = {
        "Mois": [str(p) for p in toutes_periodes],
        "AC (Dépenses réelles)": [depenses_cumulees.get(p, 0) for p in toutes_periodes],
    }

    if pv_cumulee is not None:
        data["PV (Budget prévu)"] = [pv_cumulee.get(p, 0) for p in toutes_periodes]

    if ev_cumulee is not None:
        data["EV (Valeur acquise)"] = [ev_cumulee.get(p, 0) for p in toutes_periodes]

    df_tableau = pd.DataFrame(data)

    # Calculer les écarts si toutes les valeurs sont disponibles
    if pv_cumulee is not None and ev_cumulee is not None:
        df_tableau["SV (Schedule Variance)"] = df_tableau.get("EV (Valeur acquise)", 0) - df_tableau.get(
            "PV (Budget prévu)", 0
        )
        df_tableau["CV (Cost Variance)"] = (
            df_tableau.get("EV (Valeur acquise)", 0) - df_tableau["AC (Dépenses réelles)"]
        )

    # Sauvegarder en CSV et Excel
    csv_file = f"{nom_base}.csv"
    xlsx_file = f"{nom_base}.xlsx"
    df_tableau.to_csv(csv_file, index=False)
    df_tableau.to_excel(xlsx_file, index=False)

    print("\n=== TABLEAU COMPARATIF EVM ===")
    print(df_tableau.to_string(index=False))
    print(f"\n✓ Tableau sauvegardé: {csv_file} et {xlsx_file}")

    return df_tableau


def generer_rapport_word(
    fichier_word,
    df_tableau,
    fichier_graphique,
    depenses_cumulees,
    pv_cumulee,
    ev_cumulee,
    projections,
    fichier_projections,
):
    """
    Génère un rapport Word complet avec définitions, tableau, deux graphiques (réalisé + projections) et conclusion

    Args:
        fichier_word: Nom du fichier Word à générer
        df_tableau: DataFrame avec le tableau comparatif
        fichier_graphique: Chemin du graphique du réalisé
        depenses_cumulees: Série des dépenses réelles (AC)
        pv_cumulee: Série du budget prévu (PV)
        ev_cumulee: Série de la valeur acquise (EV)
        projections: Dictionnaire des projections {'cpi': {...}, 'cpi_spi': {...}, 'reste_plan': {...}, 'forecast': {...}}
        fichier_projections: Chemin du graphique des projections
    """
    print(f"\n=== Génération du rapport Word: {fichier_word} ===")

    doc = Document()

    # Titre principal
    titre = doc.add_heading("Rapport d'Analyse EVM", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_rapport = datetime.now().astimezone().strftime("%d/%m/%Y")
    p = doc.add_paragraph(f"Date du rapport: {date_rapport}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Section 1: Définitions
    doc.add_heading("1. Définitions EVM", 1)

    definitions = [
        (
            "PV (Planned Value)",
            "Budget prévu ou valeur planifiée. Représente le coût budgété du travail prévu à une date donnée.",
        ),
        (
            "AC (Actual Cost)",
            "Coût réel ou dépenses réelles. Représente le coût réel du travail effectué à une date donnée.",
        ),
        (
            "EV (Earned Value)",
            "Valeur acquise ou valeur gagnée. Représente la valeur du travail réellement accompli à une date donnée, mesurée en termes de budget.",
        ),
        (
            "EAC (Estimate at Completion)",
            "Estimation à terminaison. Projection du coût total du projet à son achèvement.",
        ),
        ("CV (Cost Variance)", "Écart de coût. CV = EV - AC. Un CV négatif indique un dépassement de coût."),
        ("SV (Schedule Variance)", "Écart de délai. SV = EV - PV. Un SV négatif indique un retard sur le planning."),
        (
            "CPI (Cost Performance Index)",
            "Indice de performance des coûts. CPI = EV / AC. Un CPI < 1 indique un dépassement de coût.",
        ),
        (
            "SPI (Schedule Performance Index)",
            "Indice de performance des délais. SPI = EV / PV. Un SPI < 1 indique un retard.",
        ),
    ]

    for terme, definition in definitions:
        p = doc.add_paragraph(style="List Bullet")
        run_terme = p.add_run(terme + ": ")
        run_terme.bold = True
        p.add_run(definition)

    doc.add_page_break()

    # Section 2: Réalisé à date
    doc.add_heading("2. Réalisé à Date", 1)

    # 2.1 Tableau de valeurs
    doc.add_heading("2.1 Tableau des Valeurs", 2)

    # Filtrer les lignes avec des valeurs non nulles pour le tableau Word
    df_filtre = df_tableau[
        (df_tableau["AC (Dépenses réelles)"] > 0)
        | (df_tableau["PV (Budget prévu)"] > 0)
        | (df_tableau["EV (Valeur acquise)"] > 0)
    ].copy()

    # Créer le tableau Word
    if len(df_filtre) > 0:
        table = doc.add_table(rows=1, cols=len(df_filtre.columns))
        table.style = "Light Grid Accent 1"

        # En-têtes
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df_filtre.columns):
            hdr_cells[i].text = col
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Données
        for _, row in df_filtre.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                if isinstance(value, (int, float)):
                    row_cells[i].text = f"{value:,.2f}"
                else:
                    row_cells[i].text = str(value)

    # 2.2 Graphique du réalisé
    doc.add_heading("2.2 Graphique du Réalisé", 2)

    if Path(fichier_graphique).exists():
        doc.add_picture(fichier_graphique, width=Inches(6.5))
        p = doc.add_paragraph("Figure 1: Courbes du réalisé - AC, PV, EV et variances")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    # 2.3 Indicateurs actuels
    doc.add_heading("2.3 Indicateurs de Performance Actuels", 2)

    # Récupérer les dernières valeurs
    dernier_mois = depenses_cumulees.index[-1]
    ac_actuel = depenses_cumulees.iloc[-1]

    ev_actuel = 0
    if ev_cumulee is not None and len(ev_cumulee) > 0:
        ev_actuel = ev_cumulee.iloc[-1]

    pv_actuel = 0
    if pv_cumulee is not None and dernier_mois in pv_cumulee.index:
        pv_actuel = pv_cumulee[dernier_mois]

    # Calculer CV et SV actuels
    cv_actuel = ev_actuel - ac_actuel
    sv_actuel = ev_actuel - pv_actuel

    # Calculer CPI et SPI
    cpi = ev_actuel / ac_actuel if ac_actuel > 0 else 0
    spi = ev_actuel / pv_actuel if pv_actuel > 0 else 0

    p = doc.add_paragraph(f"Au mois de {dernier_mois}:")

    indicateurs_actuels = [
        ("Dépenses Réelles (AC)", f"{ac_actuel:,.2f} €"),
        ("Valeur Acquise (EV)", f"{ev_actuel:,.2f} €"),
        ("Valeur Planifiée (PV)", f"{pv_actuel:,.2f} €"),
        ("", ""),
        ("Cost Variance (CV)", f"{cv_actuel:,.2f} €", cv_actuel),
        ("Schedule Variance (SV)", f"{sv_actuel:,.2f} €", sv_actuel),
        ("Cost Performance Index (CPI)", f"{cpi:.2f}", cpi - 1),
        ("Schedule Performance Index (SPI)", f"{spi:.2f}", spi - 1),
    ]

    for item in indicateurs_actuels:
        if len(item) >= 2:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{item[0]}: {item[1]}")
            if len(item) > 2:
                valeur = item[2]
                if valeur < 0:
                    run.font.color.rgb = RGBColor(231, 76, 60)  # Rouge
                elif valeur > 0:
                    run.font.color.rgb = RGBColor(46, 204, 113)  # Vert

    # Interprétation
    doc.add_paragraph()
    p = doc.add_paragraph("Interprétation:")
    p.runs[0].bold = True

    if cv_actuel < 0:
        doc.add_paragraph(
            f"⚠ Le projet présente un dépassement de coût de {abs(cv_actuel):,.2f} € à date.", style="List Bullet"
        )
    else:
        doc.add_paragraph(
            f"✓ Le projet est sous budget avec une économie de {cv_actuel:,.2f} € à date.", style="List Bullet"
        )

    if sv_actuel < 0:
        doc.add_paragraph(
            f"⚠ Le projet présente un retard équivalent à {abs(sv_actuel):,.2f} € de travail non réalisé.",
            style="List Bullet",
        )
    else:
        doc.add_paragraph(
            f"✓ Le projet est en avance avec {sv_actuel:,.2f} € de travail supplémentaire réalisé.", style="List Bullet"
        )

    if cpi < 1:
        doc.add_paragraph(
            f"⚠ L'efficacité des coûts est de {cpi*100:.1f}% (chaque euro dépensé génère {cpi:.2f} € de valeur).",
            style="List Bullet",
        )
    else:
        doc.add_paragraph(
            f"✓ L'efficacité des coûts est de {cpi*100:.1f}% (chaque euro dépensé génère {cpi:.2f} € de valeur).",
            style="List Bullet",
        )

    doc.add_page_break()

    # Section 3: Projections à Terminaison
    doc.add_heading("3. Projections à Terminaison", 1)

    # 3.1 Tableau des scénarios
    doc.add_heading("3.1 Tableau Comparatif des Scénarios", 2)

    budget_total = 0
    if pv_cumulee is not None:
        budget_total = pv_cumulee.iloc[-1]

    scenarios_data = []

    labels_scenarios = {
        "reste_plan": "Méthode Reste à Plan (Optimiste)",
        "cpi": "Méthode CPI (Réaliste)",
        "cpi_spi": "Méthode CPI×SPI (Pessimiste)",
        "forecast": "Forecast Manuel",
    }

    for methode, label in labels_scenarios.items():
        if methode in projections and projections[methode] is not None:
            proj = projections[methode]
            eac = proj["eac"]
            date_fin = proj["date"].strftime("%m/%Y") if proj["date"] else "N/A"
            vac = budget_total - eac
            scenarios_data.append(
                {
                    "Scénario": label,
                    "EAC (€)": eac,
                    "Date Fin": date_fin,
                    "VAC (€)": vac,
                    "Dépassement": "Oui" if vac < 0 else "Non",
                }
            )

    if scenarios_data:
        df_scenarios = pd.DataFrame(scenarios_data)

        # Créer le tableau dans Word
        table = doc.add_table(rows=1, cols=len(df_scenarios.columns))
        table.style = "Light Grid Accent 1"

        # En-têtes
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df_scenarios.columns):
            hdr_cells[i].text = col
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Données
        for _, row in df_scenarios.iterrows():
            row_cells = table.add_row().cells
            for i, (col, value) in enumerate(row.items()):
                if col in ["EAC (€)", "VAC (€)"]:
                    row_cells[i].text = f"{value:,.2f}"
                    # Colorer en rouge si VAC négatif
                    if col == "VAC (€)" and value < 0:
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(231, 76, 60)
                else:
                    row_cells[i].text = str(value)

    # 3.2 Graphique des projections
    doc.add_heading("3.2 Graphique des Projections", 2)

    if Path(fichier_projections).exists():
        doc.add_picture(fichier_projections, width=Inches(6.5))
        p = doc.add_paragraph("Figure 2: Projections à terminaison - Différents scénarios EAC")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    # 3.3 Analyse des scénarios
    doc.add_heading("3.3 Analyse des Scénarios", 2)

    if scenarios_data:
        # Créer un dictionnaire des EAC par scénario avec leur label complet
        scenarios_dict = {s["Scénario"]: s["EAC (€)"] for s in scenarios_data}

        # Identifier les scénarios disponibles
        eac_optimiste = None
        eac_realiste = None
        eac_pessimiste = None
        label_optimiste = ""
        label_realiste = ""
        label_pessimiste = ""

        # Chercher chaque type de scénario
        for scenario_label, eac_val in scenarios_dict.items():
            if "Optimiste" in scenario_label:
                eac_optimiste = eac_val
                label_optimiste = scenario_label
            elif "Réaliste" in scenario_label:
                eac_realiste = eac_val
                label_realiste = scenario_label
            elif "Pessimiste" in scenario_label:
                eac_pessimiste = eac_val
                label_pessimiste = scenario_label

        # Si on a un forecast manuel, on l'analyse aussi
        eac_forecast = None
        label_forecast = ""
        for scenario_label, eac_val in scenarios_dict.items():
            if "Forecast" in scenario_label or "Manuel" in scenario_label:
                eac_forecast = eac_val
                label_forecast = scenario_label

        doc.add_paragraph(f"Budget Total (BAC): {budget_total:,.2f} €")
        doc.add_paragraph()
        doc.add_paragraph("Fourchette des projections:")

        # Afficher les scénarios disponibles
        if eac_optimiste is not None:
            doc.add_paragraph(f"  • {label_optimiste}: {eac_optimiste:,.2f} €", style="List Bullet")
        if eac_realiste is not None:
            doc.add_paragraph(f"  • {label_realiste}: {eac_realiste:,.2f} €", style="List Bullet")
        if eac_pessimiste is not None:
            doc.add_paragraph(f"  • {label_pessimiste}: {eac_pessimiste:,.2f} €", style="List Bullet")
        if eac_forecast is not None:
            doc.add_paragraph(f"  • {label_forecast}: {eac_forecast:,.2f} €", style="List Bullet")

        doc.add_paragraph()

        p = doc.add_paragraph("Écarts par rapport au budget:")
        p.runs[0].bold = True

        # Calculer les écarts pour chaque scénario disponible
        scenarios_ecarts = []
        if eac_optimiste is not None:
            scenarios_ecarts.append(("Optimiste", eac_optimiste - budget_total, eac_optimiste))
        if eac_realiste is not None:
            scenarios_ecarts.append(("Réaliste", eac_realiste - budget_total, eac_realiste))
        if eac_pessimiste is not None:
            scenarios_ecarts.append(("Pessimiste", eac_pessimiste - budget_total, eac_pessimiste))
        if eac_forecast is not None:
            scenarios_ecarts.append(("Forecast", eac_forecast - budget_total, eac_forecast))

        for label, ecart, _eac_val in scenarios_ecarts:
            p = doc.add_paragraph(style="List Bullet")
            signe = "+" if ecart >= 0 else ""
            run = p.add_run(f"  • {label}: {signe}{ecart:,.2f} € ({(ecart/budget_total)*100:+.1f}%)")
            if ecart < 0:
                run.font.color.rgb = RGBColor(231, 76, 60)  # Rouge
            else:
                run.font.color.rgb = RGBColor(46, 204, 113)  # Vert

    doc.add_page_break()

    # Section 4: Conclusion et Recommandations
    doc.add_heading("4. Conclusion et Recommandations", 1)

    doc.add_heading("4.1 Synthèse", 2)

    # Synthèse de la situation actuelle
    doc.add_paragraph("Performance actuelle:")
    if cpi < 0.9:
        doc.add_paragraph(
            "⚠ Le CPI est très faible, indiquant une efficacité des coûts préoccupante. Actions correctives urgentes recommandées.",
            style="List Bullet",
        )
    elif cpi < 1:
        doc.add_paragraph(
            "⚠ Le CPI est inférieur à 1, indiquant un dépassement de coût. Une surveillance étroite est nécessaire.",
            style="List Bullet",
        )
    else:
        doc.add_paragraph("✓ Le CPI est supérieur à 1, indiquant une bonne efficacité des coûts.", style="List Bullet")

    if spi < 0.9:
        doc.add_paragraph(
            "⚠ Le SPI est très faible, indiquant un retard significatif. Révision du planning recommandée.",
            style="List Bullet",
        )
    elif spi < 1:
        doc.add_paragraph(
            "⚠ Le SPI est inférieur à 1, indiquant un retard. Des mesures d'accélération devraient être envisagées.",
            style="List Bullet",
        )
    else:
        doc.add_paragraph(
            "✓ Le SPI est supérieur à 1, indiquant une bonne performance sur les délais.", style="List Bullet"
        )

    doc.add_paragraph()

    # Projections - Calculer min/max/tous les EAC disponibles
    if scenarios_data:
        eacs_disponibles = [
            eac_val for eac_val in [eac_optimiste, eac_realiste, eac_pessimiste, eac_forecast] if eac_val is not None
        ]

        if eacs_disponibles:
            doc.add_paragraph("Projections à terminaison:")

            if all(eac > budget_total for eac in eacs_disponibles):
                doc.add_paragraph(
                    "⚠ Tous les scénarios prévoient un dépassement de budget. Des mesures correctives sont nécessaires.",
                    style="List Bullet",
                )
            elif any(eac > budget_total for eac in eacs_disponibles):
                doc.add_paragraph(
                    "⚠ Certains scénarios prévoient un dépassement de budget. Une vigilance accrue est requise.",
                    style="List Bullet",
                )
            else:
                doc.add_paragraph(
                    "✓ Les projections indiquent un achèvement sous budget dans tous les scénarios.",
                    style="List Bullet",
                )

            # Calculer l'écart entre le meilleur et le pire scénario
            eac_min_calc = min(eacs_disponibles)
            eac_max_calc = max(eacs_disponibles)
            ecart_relatif = ((eac_max_calc - eac_min_calc) / budget_total) * 100
            if ecart_relatif > 10:
                doc.add_paragraph(
                    f"⚠ L'écart entre scénarios est important ({ecart_relatif:.1f}% du budget), reflétant une forte incertitude.",
                    style="List Bullet",
                )
            elif ecart_relatif > 5:
                doc.add_paragraph(
                    f"ℹ L'écart entre scénarios est modéré ({ecart_relatif:.1f}% du budget).", style="List Bullet"
                )
            else:
                doc.add_paragraph(
                    f"✓ L'écart entre scénarios est faible ({ecart_relatif:.1f}% du budget), indiquant une bonne prévisibilité.",
                    style="List Bullet",
                )

    doc.add_heading("4.2 Recommandations", 2)

    if cpi < 1 or spi < 1 or (scenarios_data and any(s["VAC (€)"] < 0 for s in scenarios_data)):
        doc.add_paragraph("Actions recommandées:")

        if cpi < 1:
            doc.add_paragraph(
                "1. Analyser les causes du dépassement de coût et identifier les postes problématiques",
                style="List Number",
            )
            doc.add_paragraph(
                "2. Mettre en place des mesures de réduction des coûts ou réviser le scope", style="List Number"
            )

        if spi < 1:
            doc.add_paragraph(
                "3. Revoir la planification et identifier les leviers d'accélération", style="List Number"
            )
            doc.add_paragraph("4. Augmenter les ressources si nécessaire pour rattraper le retard", style="List Number")

        if scenarios_data and any(s["VAC (€)"] < 0 for s in scenarios_data):
            doc.add_paragraph(
                "5. Prévoir un budget de contingence pour couvrir le dépassement projeté", style="List Number"
            )
            doc.add_paragraph(
                "6. Communiquer proactivement avec les parties prenantes sur les risques financiers",
                style="List Number",
            )
    else:
        doc.add_paragraph("Le projet montre de bonnes performances. Recommandations:", style="List Bullet")
        doc.add_paragraph("  • Maintenir les pratiques actuelles de gestion", style="List Bullet")
        doc.add_paragraph("  • Continuer la surveillance régulière des indicateurs", style="List Bullet")
        doc.add_paragraph("  • Capitaliser sur les bonnes pratiques pour les projets futurs", style="List Bullet")

    # Sauvegarder le document
    doc.save(fichier_word)
    print(f"✓ Rapport Word généré: {fichier_word}")


def main():
    """
    Fonction principale
    """
    # Parser les arguments
    args = parser_arguments()

    print("=== Analyse des dépenses SAP ===\n")
    print("Fichiers d'entrée:")
    print(f"  - SAP:      {args.sap}")
    print(f"  - PV:       {args.pv}")
    print(f"  - VA:       {args.va}")
    print(f"  - Forecast: {args.forecast}")
    print("\nFichiers de sortie:")
    print(f"  - Graphique: {args.output}")
    print(f"  - Tableaux:  {args.tableau}.csv / {args.tableau}.xlsx\n")

    # Lecture du fichier Excel
    df = lire_export_sap(args.sap)

    if df is None:
        return

    # Affichage d'un échantillon des données
    print("\nAperçu des données:")
    print(df.head())

    # Colonnes spécifiques pour cet export SAP
    colonne_date = "Date de la pièce"
    colonne_montant = "Val./Devise objet"

    # Vérification que les colonnes existent
    if colonne_date not in df.columns or colonne_montant not in df.columns:
        print("\n⚠️  Colonnes requises non trouvées.")
        print(f"Colonnes disponibles: {df.columns.tolist()}")
        return

    print(f"\n✓ Colonne date utilisée: {colonne_date}")
    print(f"✓ Colonne montant utilisée: {colonne_montant}")

    # Calcul des dépenses cumulées
    depenses_cumulees = calculer_depenses_cumulees(df, colonne_date, colonne_montant)

    print("\nDépenses cumulées par mois (AC):")
    print(depenses_cumulees)

    # Lecture et traitement de la Planned Value
    df_pv = lire_planned_value(args.pv)
    result = traiter_planned_value(df_pv)

    pv_cumulee = None
    jalons = None
    if result is not None:
        if isinstance(result, tuple):
            pv_cumulee, jalons = result
        else:
            pv_cumulee = result
        print("\nPlanned Value cumulée par mois (PV):")
        print(pv_cumulee)

    # Lecture et calcul de l'Earned Value
    df_va = lire_valeur_acquise(args.va)
    ev_cumulee = calculer_earned_value(df_pv, df_va)

    if ev_cumulee is not None:
        print("\nEarned Value cumulée par mois (EV):")
        print(ev_cumulee)

    # Calcul automatique des projections EAC avec différentes méthodes
    result_proj = calculer_projections_automatiques(depenses_cumulees, ev_cumulee, pv_cumulee, df_pv)

    if result_proj is not None:
        projections_data, series_projections, date_fin_proj = result_proj

        # Convertir au format attendu par les fonctions de tracé
        projections = {}
        for methode in ["CPI", "CPI_SPI", "RESTE_PLAN"]:
            if methode in projections_data and methode in series_projections:
                projections[methode.lower()] = {
                    "series": series_projections[methode],
                    "eac": projections_data[methode]["eac"],
                    "date": date_fin_proj.to_timestamp() if hasattr(date_fin_proj, "to_timestamp") else date_fin_proj,
                }
    else:
        projections = {}

    # Si un fichier forecast est fourni, l'ajouter aux projections
    df_forecast = lire_forecast(args.forecast)
    if df_forecast is not None and ev_cumulee is not None:
        result = calculer_eac_projete(ev_cumulee, df_forecast, df_pv)
        if result is not None:
            eac_projete, _ = result
            # Ajouter la projection forecast au dictionnaire
            projections["forecast"] = {
                "series": eac_projete,
                "eac": eac_projete.iloc[-1] if len(eac_projete) > 0 else 0,
                "date": eac_projete.index[-1].to_timestamp() if len(eac_projete) > 0 else None,
            }
            print("\nEAC projeté (forecast manuel) par mois:")
            print(eac_projete)

    # Générer le tableau comparatif
    df_tableau = generer_tableau_comparatif(depenses_cumulees, pv_cumulee, ev_cumulee, args.tableau)

    # Tracer les deux graphiques séparés
    fichier_realise = args.output.replace(".png", "_realise.png")
    fichier_projections = args.output.replace(".png", "_projections.png")

    tracer_courbe_realise(depenses_cumulees, pv_cumulee, jalons, ev_cumulee, fichier_realise)
    tracer_courbe_projections(depenses_cumulees, ev_cumulee, projections, fichier_projections)

    # Générer le rapport Word si demandé
    if args.word:
        generer_rapport_word(
            args.word,
            df_tableau,
            fichier_realise,
            depenses_cumulees,
            pv_cumulee,
            ev_cumulee,
            projections,
            fichier_projections,
        )

        # Supprimer les fichiers intermédiaires
        print("\n=== Nettoyage des fichiers intermédiaires ===")
        fichiers_a_supprimer = [fichier_realise, fichier_projections, f"{args.tableau}.csv", f"{args.tableau}.xlsx"]

        for fichier in fichiers_a_supprimer:
            fichier_path = Path(fichier)
            if fichier_path.exists():
                fichier_path.unlink()
                print(f"✓ {fichier} supprimé")

    print("\n✓ Analyse terminée avec succès!")


if __name__ == "__main__":
    main()
