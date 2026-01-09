"""
Module de génération de rapports Word
"""

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor

LIST_BULLET_STYLE = "List Bullet"
LIST_NUMBER_STYLE = "List Number"

COL_AC = "AC (Dépenses réelles)"
COL_PV = "PV (Budget prévu)"
COL_EV = "EV (Valeur acquise)"
COL_EAC = "EAC (€)"
COL_VAC = "VAC (€)"

COLOR_RED = RGBColor(231, 76, 60)
COLOR_GREEN = RGBColor(46, 204, 113)


def _add_centered_paragraph(doc: DocumentType, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_title_page(doc: DocumentType) -> None:
    titre = doc.add_heading("Rapport d'Analyse EVM", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_rapport = datetime.now().astimezone().strftime("%d/%m/%Y")
    _add_centered_paragraph(doc, f"Date du rapport: {date_rapport}")
    doc.add_page_break()


def _add_definitions_section(doc: DocumentType) -> None:
    doc.add_heading("1. Définitions EVM", 1)

    definitions = [
        (
            "PV (Planned Value)",
            "Budget prévu ou valeur planifiée. Représente le coût budgété du travail prévu à une date donnée.",
        ),
        (
            "AC (Actual Cost)",
            "Coût réel ou dépenses réelles. Représente le coût réel du travail effectué à une date donnée.",
        ),
        (
            "EV (Earned Value)",
            "Valeur acquise ou valeur gagnée. Représente la valeur du travail réellement accompli à une date donnée, mesurée en termes de budget.",
        ),
        (
            "EAC (Estimate at Completion)",
            "Estimation à terminaison. Projection du coût total du projet à son achèvement.",
        ),
        ("CV (Cost Variance)", "Écart de coût. CV = EV - AC. Un CV négatif indique un dépassement de coût."),
        ("SV (Schedule Variance)", "Écart de délai. SV = EV - PV. Un SV négatif indique un retard sur le planning."),
        (
            "CPI (Cost Performance Index)",
            "Indice de performance des coûts. CPI = EV / AC. Un CPI < 1 indique un dépassement de coût.",
        ),
        (
            "SPI (Schedule Performance Index)",
            "Indice de performance des délais. SPI = EV / PV. Un SPI < 1 indique un retard.",
        ),
    ]

    for terme, definition in definitions:
        p = doc.add_paragraph(style=LIST_BULLET_STYLE)
        run_terme = p.add_run(terme + ": ")
        run_terme.bold = True
        p.add_run(definition)

    doc.add_page_break()


def _filter_df_tableau(df_tableau: pd.DataFrame) -> pd.DataFrame:
    conditions = [df_tableau[COL_AC] > 0]
    if COL_PV in df_tableau.columns:
        conditions.append(df_tableau[COL_PV] > 0)
    if COL_EV in df_tableau.columns:
        conditions.append(df_tableau[COL_EV] > 0)

    combined = conditions[0]
    for cond in conditions[1:]:
        combined = combined | cond

    return df_tableau[combined].copy()


def _add_dataframe_table(doc: DocumentType, df: pd.DataFrame) -> None:
    if len(df) == 0:
        return

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = f"{value:,.2f}" if isinstance(value, (int, float)) else str(value)


def _add_picture_with_caption(doc: DocumentType, path: str, caption: str) -> None:
    if not Path(path).exists():
        return
    doc.add_picture(path, width=Inches(6.5))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True


def _compute_current_metrics(depenses_cumulees, pv_cumulee, ev_cumulee):
    dernier_mois = depenses_cumulees.index[-1]
    ac_actuel = depenses_cumulees.iloc[-1]

    ev_actuel = ev_cumulee.iloc[-1] if ev_cumulee is not None and len(ev_cumulee) > 0 else 0
    pv_actuel = pv_cumulee[dernier_mois] if pv_cumulee is not None and dernier_mois in pv_cumulee.index else 0

    cv_actuel = ev_actuel - ac_actuel
    sv_actuel = ev_actuel - pv_actuel
    cpi = ev_actuel / ac_actuel if ac_actuel > 0 else 0
    spi = ev_actuel / pv_actuel if pv_actuel > 0 else 0

    return {
        "dernier_mois": dernier_mois,
        "ac": ac_actuel,
        "ev": ev_actuel,
        "pv": pv_actuel,
        "cv": cv_actuel,
        "sv": sv_actuel,
        "cpi": cpi,
        "spi": spi,
    }


def _add_colored_bullet(doc: DocumentType, text: str, delta: float | None = None) -> None:
    p = doc.add_paragraph(style=LIST_BULLET_STYLE)
    run = p.add_run(text)
    if delta is None:
        return
    if delta < 0:
        run.font.color.rgb = COLOR_RED
    elif delta > 0:
        run.font.color.rgb = COLOR_GREEN


def _add_realise_section(
    doc: DocumentType,
    df_tableau: pd.DataFrame,
    fichier_graphique: str,
    depenses_cumulees,
    pv_cumulee,
    ev_cumulee,
):
    doc.add_heading("2. Réalisé à Date", 1)
    doc.add_heading("2.1 Tableau des Valeurs", 2)
    df_filtre = _filter_df_tableau(df_tableau)
    _add_dataframe_table(doc, df_filtre)

    doc.add_heading("2.2 Graphique du Réalisé", 2)
    _add_picture_with_caption(doc, fichier_graphique, "Figure 1: Courbes du réalisé - AC, PV, EV et variances")

    doc.add_heading("2.3 Indicateurs de Performance Actuels", 2)
    metrics = _compute_current_metrics(depenses_cumulees, pv_cumulee, ev_cumulee)
    doc.add_paragraph(f"Au mois de {metrics['dernier_mois']}:")

    _add_colored_bullet(doc, f"Dépenses Réelles (AC): {metrics['ac']:,.2f} €")
    _add_colored_bullet(doc, f"Valeur Acquise (EV): {metrics['ev']:,.2f} €")
    _add_colored_bullet(doc, f"Valeur Planifiée (PV): {metrics['pv']:,.2f} €")
    doc.add_paragraph()
    _add_colored_bullet(doc, f"Cost Variance (CV): {metrics['cv']:,.2f} €", metrics["cv"])
    _add_colored_bullet(doc, f"Schedule Variance (SV): {metrics['sv']:,.2f} €", metrics["sv"])
    _add_colored_bullet(doc, f"Cost Performance Index (CPI): {metrics['cpi']:.2f}", metrics["cpi"] - 1)
    _add_colored_bullet(doc, f"Schedule Performance Index (SPI): {metrics['spi']:.2f}", metrics["spi"] - 1)

    doc.add_paragraph()
    p = doc.add_paragraph("Interprétation:")
    p.runs[0].bold = True

    cv = metrics["cv"]
    sv = metrics["sv"]
    cpi = metrics["cpi"]
    if cv < 0:
        doc.add_paragraph(
            f"⚠ Le projet présente un dépassement de coût de {abs(cv):,.2f} € à date.", style=LIST_BULLET_STYLE
        )
    else:
        doc.add_paragraph(
            f"✓ Le projet est sous budget avec une économie de {cv:,.2f} € à date.", style=LIST_BULLET_STYLE
        )

    if sv < 0:
        doc.add_paragraph(
            f"⚠ Le projet présente un retard équivalent à {abs(sv):,.2f} € de travail non réalisé.",
            style=LIST_BULLET_STYLE,
        )
    else:
        doc.add_paragraph(
            f"✓ Le projet est en avance avec {sv:,.2f} € de travail supplémentaire réalisé.", style=LIST_BULLET_STYLE
        )

    if cpi < 1:
        doc.add_paragraph(
            f"⚠ L'efficacité des coûts est de {cpi * 100:.1f}% (chaque euro dépensé génère {cpi:.2f} € de valeur).",
            style=LIST_BULLET_STYLE,
        )
    else:
        doc.add_paragraph(
            f"✓ L'efficacité des coûts est de {cpi * 100:.1f}% (chaque euro dépensé génère {cpi:.2f} € de valeur).",
            style=LIST_BULLET_STYLE,
        )

    doc.add_page_break()
    return metrics


def _build_scenarios_data(projections, budget_total: float):
    labels_scenarios = {
        "reste_plan": "Méthode Reste à Plan (Optimiste)",
        "cpi": "Méthode CPI (Réaliste)",
        "cpi_spi": "Méthode CPI×SPI (Pessimiste)",
        "forecast": "Forecast Manuel",
    }
    scenarios_data: list[dict] = []
    for methode, label in labels_scenarios.items():
        if methode in projections and projections[methode] is not None:
            proj = projections[methode]
            eac = proj["eac"]
            date_fin = proj["date"].strftime("%m/%Y") if proj["date"] else "N/A"
            vac = budget_total - eac
            scenarios_data.append(
                {
                    "Scénario": label,
                    COL_EAC: eac,
                    "Date Fin": date_fin,
                    COL_VAC: vac,
                    "Dépassement": "Oui" if vac < 0 else "Non",
                }
            )
    return scenarios_data


def _bold_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True


def _color_cell(cell, color: RGBColor) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color


def _fill_scenario_row(row_cells, row: pd.Series) -> None:
    for i, (col, value) in enumerate(row.items()):
        if col in [COL_EAC, COL_VAC]:
            row_cells[i].text = f"{value:,.2f}"
            if col == COL_VAC and value < 0:
                _color_cell(row_cells[i], COLOR_RED)
        else:
            row_cells[i].text = str(value)


def _add_scenarios_table(doc: DocumentType, scenarios_data: list[dict]) -> None:
    if not scenarios_data:
        return

    df_scenarios = pd.DataFrame(scenarios_data)
    table = doc.add_table(rows=1, cols=len(df_scenarios.columns))
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_scenarios.columns):
        hdr_cells[i].text = col
        _bold_cell(hdr_cells[i])

    for _, row in df_scenarios.iterrows():
        row_cells = table.add_row().cells
        _fill_scenario_row(row_cells, row)


def _add_projections_section(doc: DocumentType, projections, fichier_projections: str, budget_total: float):
    doc.add_heading("3. Projections à Terminaison", 1)
    doc.add_heading("3.1 Tableau Comparatif des Scénarios", 2)
    scenarios_data = _build_scenarios_data(projections, budget_total)
    _add_scenarios_table(doc, scenarios_data)

    doc.add_heading("3.2 Graphique des Projections", 2)
    _add_picture_with_caption(
        doc, fichier_projections, "Figure 2: Projections à terminaison - Différents scénarios EAC"
    )

    doc.add_heading("3.3 Analyse des Scénarios", 2)
    if scenarios_data:
        scenarios_dict = {s["Scénario"]: s[COL_EAC] for s in scenarios_data}

        eac_optimiste = next((v for k, v in scenarios_dict.items() if "Optimiste" in k), None)
        eac_realiste = next((v for k, v in scenarios_dict.items() if "Réaliste" in k), None)
        eac_pessimiste = next((v for k, v in scenarios_dict.items() if "Pessimiste" in k), None)
        eac_forecast = next((v for k, v in scenarios_dict.items() if ("Forecast" in k or "Manuel" in k)), None)

        doc.add_paragraph(f"Budget Total (BAC): {budget_total:,.2f} €")
        doc.add_paragraph()
        doc.add_paragraph("Fourchette des projections:")
        for label, eac_val in scenarios_dict.items():
            doc.add_paragraph(f"  • {label}: {eac_val:,.2f} €", style=LIST_BULLET_STYLE)

        doc.add_paragraph()
        p = doc.add_paragraph("Écarts par rapport au budget:")
        p.runs[0].bold = True

        for label, eac_val in [
            ("Optimiste", eac_optimiste),
            ("Réaliste", eac_realiste),
            ("Pessimiste", eac_pessimiste),
            ("Forecast", eac_forecast),
        ]:
            if eac_val is None:
                continue
            ecart = eac_val - budget_total
            signe = "+" if ecart >= 0 else ""
            p = doc.add_paragraph(style=LIST_BULLET_STYLE)
            run = p.add_run(f"  • {label}: {signe}{ecart:,.2f} € ({(ecart / budget_total) * 100:+.1f}%)")
            run.font.color.rgb = COLOR_GREEN if ecart >= 0 else COLOR_RED

    doc.add_page_break()
    return scenarios_data


def _add_performance_summary(doc: DocumentType, *, cpi: float, spi: float) -> None:
    doc.add_paragraph("Performance actuelle:")
    if cpi < 0.9:
        doc.add_paragraph(
            "⚠ Le CPI est très faible, indiquant une efficacité des coûts préoccupante. Actions correctives urgentes recommandées.",
            style=LIST_BULLET_STYLE,
        )
    elif cpi < 1:
        doc.add_paragraph(
            "⚠ Le CPI est inférieur à 1, indiquant un dépassement de coût. Une surveillance étroite est nécessaire.",
            style=LIST_BULLET_STYLE,
        )
    else:
        doc.add_paragraph(
            "✓ Le CPI est supérieur à 1, indiquant une bonne efficacité des coûts.",
            style=LIST_BULLET_STYLE,
        )

    if spi < 0.9:
        doc.add_paragraph(
            "⚠ Le SPI est très faible, indiquant un retard significatif. Révision du planning recommandée.",
            style=LIST_BULLET_STYLE,
        )
    elif spi < 1:
        doc.add_paragraph(
            "⚠ Le SPI est inférieur à 1, indiquant un retard. Des mesures d'accélération devraient être envisagées.",
            style=LIST_BULLET_STYLE,
        )
    else:
        doc.add_paragraph(
            "✓ Le SPI est supérieur à 1, indiquant une bonne performance sur les délais.",
            style=LIST_BULLET_STYLE,
        )

    doc.add_paragraph()


def _add_projection_summary(doc: DocumentType, scenarios_data: list[dict], budget_total: float) -> None:
    if not scenarios_data:
        return

    eacs_disponibles = [s[COL_EAC] for s in scenarios_data]
    if not eacs_disponibles:
        return

    doc.add_paragraph("Projections à terminaison:")
    all_over = all(eac > budget_total for eac in eacs_disponibles)
    any_over = any(eac > budget_total for eac in eacs_disponibles)
    if all_over:
        doc.add_paragraph(
            "⚠ Tous les scénarios prévoient un dépassement de budget. Des mesures correctives sont nécessaires.",
            style=LIST_BULLET_STYLE,
        )
    elif any_over:
        doc.add_paragraph(
            "⚠ Certains scénarios prévoient un dépassement de budget. Une vigilance accrue est requise.",
            style=LIST_BULLET_STYLE,
        )
    else:
        doc.add_paragraph(
            "✓ Les projections indiquent un achèvement sous budget dans tous les scénarios.",
            style=LIST_BULLET_STYLE,
        )

    eac_min_calc = min(eacs_disponibles)
    eac_max_calc = max(eacs_disponibles)
    ecart_relatif = ((eac_max_calc - eac_min_calc) / budget_total) * 100
    if ecart_relatif > 10:
        doc.add_paragraph(
            f"⚠ L'écart entre scénarios est important ({ecart_relatif:.1f}% du budget), reflétant une forte incertitude.",
            style=LIST_BULLET_STYLE,
        )
    elif ecart_relatif > 5:
        doc.add_paragraph(
            f"ℹ L'écart entre scénarios est modéré ({ecart_relatif:.1f}% du budget).",
            style=LIST_BULLET_STYLE,
        )
    else:
        doc.add_paragraph(
            f"✓ L'écart entre scénarios est faible ({ecart_relatif:.1f}% du budget), indiquant une bonne prévisibilité.",
            style=LIST_BULLET_STYLE,
        )


def _add_recommendations(doc: DocumentType, *, cpi: float, spi: float, has_negative_vac: bool) -> None:
    doc.add_heading("4.2 Recommandations", 2)

    if cpi < 1 or spi < 1 or has_negative_vac:
        doc.add_paragraph("Actions recommandées:")

        if cpi < 1:
            doc.add_paragraph(
                "1. Analyser les causes du dépassement de coût et identifier les postes problématiques",
                style=LIST_NUMBER_STYLE,
            )
            doc.add_paragraph(
                "2. Mettre en place des mesures de réduction des coûts ou réviser le scope",
                style=LIST_NUMBER_STYLE,
            )

        if spi < 1:
            doc.add_paragraph(
                "3. Revoir la planification et identifier les leviers d'accélération",
                style=LIST_NUMBER_STYLE,
            )
            doc.add_paragraph(
                "4. Augmenter les ressources si nécessaire pour rattraper le retard",
                style=LIST_NUMBER_STYLE,
            )

        if has_negative_vac:
            doc.add_paragraph(
                "5. Prévoir un budget de contingence pour couvrir le dépassement projeté",
                style=LIST_NUMBER_STYLE,
            )
            doc.add_paragraph(
                "6. Communiquer proactivement avec les parties prenantes sur les risques financiers",
                style=LIST_NUMBER_STYLE,
            )
        return

    doc.add_paragraph("Le projet montre de bonnes performances. Recommandations:", style=LIST_BULLET_STYLE)
    doc.add_paragraph("  • Maintenir les pratiques actuelles de gestion", style=LIST_BULLET_STYLE)
    doc.add_paragraph("  • Continuer la surveillance régulière des indicateurs", style=LIST_BULLET_STYLE)
    doc.add_paragraph("  • Capitaliser sur les bonnes pratiques pour les projets futurs", style=LIST_BULLET_STYLE)


def _add_conclusion_section(
    doc: DocumentType, *, cpi: float, spi: float, scenarios_data: list[dict], budget_total: float
) -> None:
    doc.add_heading("4. Conclusion et Recommandations", 1)
    doc.add_heading("4.1 Synthèse", 2)

    _add_performance_summary(doc, cpi=cpi, spi=spi)
    _add_projection_summary(doc, scenarios_data, budget_total)
    has_negative_vac = any(s[COL_VAC] < 0 for s in scenarios_data) if scenarios_data else False
    _add_recommendations(doc, cpi=cpi, spi=spi, has_negative_vac=has_negative_vac)


def generer_rapport_word(
    fichier_word,
    df_tableau,
    fichier_graphique,
    depenses_cumulees,
    pv_cumulee,
    ev_cumulee,
    projections,
    fichier_projections,
):
    """
    Génère un rapport Word complet avec définitions, tableau, deux graphiques (réalisé + projections) et conclusion

    Args:
        fichier_word: Nom du fichier Word à générer
        df_tableau: DataFrame avec le tableau comparatif
        fichier_graphique: Chemin du graphique du réalisé
        depenses_cumulees: Série des dépenses réelles (AC)
        pv_cumulee: Série du budget prévu (PV)
        ev_cumulee: Série de la valeur acquise (EV)
        projections: Dictionnaire des projections {'cpi': {...}, 'cpi_spi': {...}, 'reste_plan': {...}, 'forecast': {...}}
        fichier_projections: Chemin du graphique des projections
    """
    print(f"\n=== Génération du rapport Word: {fichier_word} ===")

    doc = Document()
    _add_title_page(doc)
    _add_definitions_section(doc)

    metrics = _add_realise_section(doc, df_tableau, fichier_graphique, depenses_cumulees, pv_cumulee, ev_cumulee)
    budget_total = pv_cumulee.iloc[-1] if pv_cumulee is not None else 0
    scenarios_data = _add_projections_section(doc, projections, fichier_projections, budget_total)
    _add_conclusion_section(
        doc,
        cpi=metrics["cpi"],
        spi=metrics["spi"],
        scenarios_data=scenarios_data,
        budget_total=budget_total,
    )

    doc.save(fichier_word)
    print(f"✓ Rapport Word généré: {fichier_word}")
